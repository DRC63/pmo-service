"""P3MAI-branded python-docx helpers shared by the Method Map document generators.

Brand: navy #0B2545 (primary), gold #C9A227 (accent), Segoe UI (reliable on the
Windows target). Produces title page, document-control table, auto TOC field,
numbered headings, captioned tables, shaded callout boxes, code blocks and a
running footer (classification | version | page).
"""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm, Twips

NAVY = RGBColor(0x0B, 0x25, 0x45)
NAVY_LIGHT = RGBColor(0x1B, 0x3F, 0x6E)
GOLD = RGBColor(0xC9, 0xA2, 0x27)
GOLD_DARK = RGBColor(0xA8, 0x84, 0x1C)
GREEN = RGBColor(0x2E, 0x7D, 0x5B)
RED = RGBColor(0xC0, 0x39, 0x2B)
PURPLE = RGBColor(0x8E, 0x5B, 0xE0)
GREY = RGBColor(0x5B, 0x66, 0x75)
TEXT = RGBColor(0x1C, 0x2B, 0x3A)

HEAD_FONT = "Segoe UI Semibold"
BODY_FONT = "Segoe UI"
CODE_FONT = "Consolas"

CALLOUTS = {
    "pitfall": ("⚠  ", "FDEEEE", RED),
    "tip": ("★  ", "FFF9E0", GOLD_DARK),
    "note": ("🔒  ", "EAF2FB", NAVY),
    "info": ("ℹ  ", "EEF1F5", NAVY_LIGHT),
}


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_border(cell, edge, hexcolor, sz=24):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:color"), hexcolor)
    borders.append(el)


def _set_font(run, name=BODY_FONT, size=11, color=TEXT, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    # ensure east-asian/complex also use the font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)


def _runs(paragraph, text, base_size=11, color=TEXT):
    """Add runs to a paragraph, honouring **bold** and `code` inline markers."""
    import re

    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            _set_font(r, BODY_FONT, base_size, color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            _set_font(r, CODE_FONT, base_size - 0.5, NAVY_LIGHT)
        else:
            r = paragraph.add_run(part)
            _set_font(r, BODY_FONT, base_size, color)


def new_doc():
    doc = Document()
    # base Normal style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    # A4, sensible margins
    for s in doc.sections:
        s.page_width = Twips(11906)
        s.page_height = Twips(16838)
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
    return doc


def _field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = ""
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(it); run._r.append(sep); run._r.append(t); run._r.append(e)


def footer(doc, classification, version):
    sec = doc.sections[0]
    f = sec.footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.text = ""
    # tab stops: left already; add center + right
    tabs = p.paragraph_format.tab_stops
    from docx.enum.text import WD_TAB_ALIGNMENT
    usable = sec.page_width - sec.left_margin - sec.right_margin
    tabs.add_tab_stop(int(usable / 2), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(int(usable), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run("P3MAI Method Map"); _set_font(r, BODY_FONT, 8, GREY)
    r = p.add_run("\t" + version + "\t"); _set_font(r, BODY_FONT, 8, GREY)
    r = p.add_run("Page "); _set_font(r, BODY_FONT, 8, GREY)
    _field(p, "PAGE");
    for run in p.runs:
        if run.text == "" and run._r.findall(qn("w:fldChar")):
            _set_font(run, BODY_FONT, 8, GREY)
    r = p.add_run("   "); _set_font(r, BODY_FONT, 8, GREY)
    r = p.add_run("| " + classification); _set_font(r, BODY_FONT, 8, GOLD_DARK, bold=True)


def title_page(doc, doc_id, title, subtitle, version, date, author, classification):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("P3MAI"); _set_font(r, HEAD_FONT, 30, NAVY, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("METHOD MAP"); _set_font(r, HEAD_FONT, 14, GOLD_DARK, bold=True)
    r.font.spacing = Pt(2)
    doc.add_paragraph()
    _rule(doc, NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(title); _set_font(r, HEAD_FONT, 26, NAVY, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle); _set_font(r, BODY_FONT, 13, GREY)
    doc.add_paragraph()
    _rule(doc, GOLD)
    for _ in range(6):
        doc.add_paragraph()
    meta = [("Document", f"{doc_id} — {title}"), ("Version", version),
            ("Date", date), ("Author", author),
            ("Classification", classification), ("Applies to", "P3MAI Method Map (PRINCE2 7)")]
    t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for k, v in meta:
        row = t.add_row().cells
        row[0].width = Cm(4.5); row[1].width = Cm(9)
        _runs(row[0].paragraphs[0], f"**{k}**")
        _runs(row[1].paragraphs[0], v)
    page_break(doc)


def _rule(doc, color):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pbdr.append(bottom); pPr.append(pbdr)


def doc_control(doc, rows):
    heading(doc, "Document control", 1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(("Version", "Date", "Author", "Change")):
        _runs(hdr[i].paragraphs[0], f"**{h}**"); _shade(hdr[i], "0B2545")
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            _runs(cells[i].paragraphs[0], val, base_size=10)
    doc.add_paragraph()


def add_toc(doc):
    heading(doc, "Contents", 1)
    p = doc.add_paragraph()
    _field(p, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    r = note.add_run("(If the contents list is blank, click it and press F9 to update.)")
    _set_font(r, BODY_FONT, 8.5, GREY, italic=True)
    page_break(doc)


_HSIZE = {1: 16, 2: 13, 3: 11.5}


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_font(r, HEAD_FONT, _HSIZE.get(level, 11), NAVY if level == 1 else NAVY_LIGHT, bold=True)
    return p


def para(doc, text, size=11):
    p = doc.add_paragraph()
    _runs(p, text, base_size=size)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    _runs(p, text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    _runs(p, text)
    return p


def figure_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); _set_font(r, BODY_FONT, 9, GREY, italic=True)


def figure(doc, path, caption, width_cm=15.5):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    figure_caption(doc, caption)
    doc.add_paragraph()


def table(doc, headers, rows, caption=None, col_widths=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption); _set_font(r, BODY_FONT, 9, GREY, italic=True)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        _shade(hdr[i], "0B2545")
        pr = hdr[i].paragraphs[0]
        r = pr.add_run(h); _set_font(r, HEAD_FONT, 10, RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if ri % 2 == 1:
                _shade(cells[i], "F3F6FA")
            _runs(cells[i].paragraphs[0], str(val), base_size=9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def callout(doc, kind, title, lines):
    prefix, fill, color = CALLOUTS[kind]
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    _cell_border(cell, "left", "%02X%02X%02X" % (color[0], color[1], color[2]), sz=36)
    p = cell.paragraphs[0]
    r = p.add_run(prefix + title)
    _set_font(r, HEAD_FONT, 10.5, color, bold=True)
    for ln in lines:
        pp = cell.add_paragraph()
        _runs(pp, ln, base_size=10)
    doc.add_paragraph()


def code_block(doc, text):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    _shade(cell, "F5F6F8")
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line if line else " ")
        _set_font(r, CODE_FONT, 9, NAVY_LIGHT)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
